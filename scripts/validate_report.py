"""Validate final Project B report outputs after layout/citation repair."""
from __future__ import annotations

import hashlib
import re
import subprocess
import zipfile
from pathlib import Path

import pandas as pd
from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "report"
DOCX = REPORT / "report.docx"
PDF = REPORT / "report.pdf"
FINAL_MD = REPORT / "report_final.md"
ASSETS = REPORT / "assets"
TABLES = ROOT / "results" / "tables"
DATA = ROOT / "results" / "data"

TITLE = "Signal Mosaic: Systematic Multi-Asset Funds with Coverage-Aware News Sentiment"
AUTHOR = "Chenhang Huang"
ZID = "z5641844"

EXPECTED_HASHES = {
    "results/data/fund_returns.csv": "5e51f0aa2044f11b181f6f26271342dad338f52e34c1d49aaebfdee408797d5c",
    "results/data/fund_weights.csv": "63d559844fa5bac5ea2f7c5af966a1bccd1d2479f2e2f1fab7a1596dd2058b26",
    "results/data/sector_sentiment_index.csv": "7df0d114663fbaf1e9721d3912fce18df0d2d201d97b27688d3e2e737a6d7f6c",
    "results/data/fusion_rebalance_signals.csv": "a5a57c00198d3cd3d7fce412a884302451633636f78581920b3af1231b6cc751",
    "results/tables/performance_metrics.csv": "bc82cf0b987675618ccd165775d112a24be071b4fb3403d848117d312c0499d8",
    "results/tables/fund_backtest_design.csv": "c2aac703e52e17e3050829ed145bfb0dba21ae7d201ace4594eca79d31725897",
    "results/tables/fund_fact_sheet_summary.csv": "f140f5e54b2fbfd29273740728bc55cb1d0db5968bec05ffcd9c256ff4c6a87b",
    "results/tables/fund_latest_holdings.csv": "8eab12b49b667f5de103963d6cfb5ae6fd78352bc3bde3c0cabb391881dddfb5",
    "results/tables/fund_optimizer_diagnostics.csv": "6347985e11b239dd73ec4d01710b4e6fda7dc1072014a063510c68ebb42544a6",
    "results/tables/fusion_before_after.csv": "d73c724a6dd073b8dccc43d3724c08101528b115b3b32baa8e25f98938142f5d",
    "results/tables/fusion_signal_diagnostics.csv": "6c96639cce5c9a40b27f1e8c86f72ada0ab933cd99ae4f739a7b64aa58ff487a",
    "results/tables/fusion_predictive_diagnostics.csv": "f37a81d5d0f7e2d5ae36fb2a564ccf31696a3fd84b76ccb8f1c46617045423de",
    "results/tables/sentiment_model_diagnostics.csv": "e4f1af4f18a3708973e0c7a5fe6bad07f8631d901d2a294d5c257d5a17c312a8",
    "results/tables/sentiment_sector_summary.csv": "eb50be195d3e583598cd9fb2e2bd7a7266997e7c6af00afe1a8e524d47d52e7f",
    "results/tables/app_artifact_inventory.csv": "97d5ff746bb1394892c03edbc636d816a40788ccdf73289dd29c426dce567fb1",
    "results/tables/pipeline_validation.csv": "f77e240eece810b0bf8acaef88ab68900a9e25f9f142645375ac289210864d4a",
}

FUNDS = [
    "Equity Equal Weight",
    "Equity Minimum Variance",
    "Equity Risk Parity",
    "Crypto Equal Weight",
    "Crypto Minimum Variance",
    "Crypto Risk Parity",
    "Combined Equal Weight",
    "Combined Minimum Variance",
    "Combined Risk Parity",
    "Equity Naive Sentiment Tilt",
    "Equity Coverage-Gated Sentiment Tilt",
]

REFERENCES = [
    "Hric, J. and Lin, Y. (2026)",
    "Hutto, C.J. and Gilbert, E.E. (2014)",
    "Markowitz, H. (1952)",
    "Sharpe, W.F. (1966)",
    "UNSW Business School (2026)",
]

CITATIONS = [
    "(Hric and Lin, 2026)",
    "(Hutto and Gilbert, 2014)",
    "(Markowitz, 1952)",
    "(Sharpe, 1966)",
    "(UNSW Business School, 2026)",
]


def document_text(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def media_count(docx_path: Path) -> int:
    with zipfile.ZipFile(docx_path) as zf:
        return len([name for name in zf.namelist() if name.startswith("word/media/")])


def require(condition: bool, message: str, failures: list[str]) -> None:
    if not condition:
        failures.append(message)


def word_page_stats() -> dict[str, int]:
    ps = f"""
$ErrorActionPreference = 'Stop'
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$word.DisplayAlerts = 0
$doc = $word.Documents.Open('{DOCX.resolve()}')
$doc.Repaginate()
$pages = $doc.ComputeStatistics(2)
$exec = 0; $conclusion = 0; $refs = 0; $apps = 0
foreach ($p in $doc.Paragraphs) {{
  $t = $p.Range.Text.Trim()
  if ($t -eq 'Executive Summary' -and $exec -eq 0) {{ $exec = $p.Range.Information(3) }}
  if ($t -eq '7. Conclusion' -and $conclusion -eq 0) {{ $conclusion = $p.Range.Information(3) }}
  if ($t -eq 'References' -and $refs -eq 0) {{ $refs = $p.Range.Information(3) }}
  if ($t -eq 'Appendices' -and $apps -eq 0) {{ $apps = $p.Range.Information(3) }}
}}
$lastNarr = $refs - 1
$lastNarrChars = 0
foreach ($p in $doc.Paragraphs) {{
  if ($p.Range.Information(3) -eq $lastNarr) {{
    $t = $p.Range.Text.Trim()
    if ($t.Length -gt 0) {{ $lastNarrChars += $t.Length }}
  }}
}}
$doc.Close($false)
$word.Quit()
Write-Output "pages=$pages"
Write-Output "exec=$exec"
Write-Output "conclusion=$conclusion"
Write-Output "refs=$refs"
Write-Output "apps=$apps"
Write-Output "last_narr_chars=$lastNarrChars"
"""
    completed = subprocess.run(
        ["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-Command", ps],
        cwd=ROOT,
        text=True,
        capture_output=True,
        check=True,
    )
    out: dict[str, int] = {}
    for line in completed.stdout.splitlines():
        if "=" in line:
            key, value = line.split("=", 1)
            out[key] = int(value)
    return out


def count_page_fields(doc: Document) -> int:
    count = 0
    for section in doc.sections:
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            xml = footer._element.xml
            count += xml.count("PAGE")
    return count


def section_page_numbering_ok(doc: Document, failures: list[str]) -> None:
    require(len(doc.sections) >= 3, "expected cover, narrative, and appendix sections", failures)
    if len(doc.sections) < 3:
        return
    cover, body, appendix = doc.sections[0], doc.sections[1], doc.sections[-1]
    require(cover.different_first_page_header_footer, "cover does not suppress first-page header/footer", failures)
    require(not body.different_first_page_header_footer, "Executive Summary page may suppress numbering", failures)
    require(not appendix.different_first_page_header_footer, "appendix first page may suppress numbering", failures)
    require("PAGE" not in cover.first_page_footer._element.xml, "cover first-page footer contains a PAGE field", failures)
    body_pg_num = body._sectPr.find(qn("w:pgNumType"))
    require(body_pg_num is not None and body_pg_num.get(qn("w:start")) == "1", "narrative numbering does not restart at 1", failures)
    for idx, section in enumerate(doc.sections[2:], start=3):
        pg_num = section._sectPr.find(qn("w:pgNumType"))
        require(pg_num is None or pg_num.get(qn("w:start")) is None, f"section {idx} restarts page numbering", failures)


def figure1_key_ok(failures: list[str]) -> None:
    source = (ROOT / "scripts" / "build_report.py").read_text(encoding="utf-8")
    required_labels = [
        "Eq EW",
        "Eq Min Var",
        "Eq Risk Parity",
        "Eq Naive Tilt",
        "Eq Coverage-Gated",
        "Crypto EW",
        "Crypto Min Var",
        "Crypto Risk Parity",
        "Combined EW",
        "Combined Min Var",
        "Combined Risk Parity",
    ]
    require((ASSETS / "fund_risk_return_all_11.png").exists(), "report-only Figure 1 asset missing", failures)
    require((ASSETS / "fund_risk_return_all_11.png").stat().st_size > 50_000, "report-only Figure 1 asset appears too small", failures)
    require("Point key" in source, "Figure 1 does not use the compact point-key design", failures)
    for label in required_labels:
        require(label in source, f"Figure 1 compact label missing from generator: {label}", failures)


def validate_references(text: str, doc: Document, failures: list[str]) -> None:
    refs_start = text.find("References")
    refs_end = text.find("Appendices")
    refs_text = text[refs_start:refs_end] if refs_start >= 0 and refs_end > refs_start else ""
    for ref in REFERENCES:
        require(ref in refs_text, f"reference missing: {ref}", failures)
    require("UNSW FINS3645 (2026a)" not in text, "Project Brief reference remains", failures)
    require("UNSW FINS3645 (2026b)" not in text, "old dataset reference remains", failures)
    require("2026a" not in text and "2026b" not in text, "2026a/2026b suffix remains", failures)
    require("Project Brief" not in refs_text, "Project Brief remains in references", failures)
    for citation in CITATIONS:
        require(citation in text, f"in-text citation missing: {citation}", failures)
    require("UNSW Business School (2026)" in refs_text, "dataset reference missing", failures)
    require("(UNSW Business School, 2026)" in text, "dataset in-text citation missing", failures)
    require(text.count("Hric, J. and Lin, Y. (2026)") == 1, "Hric reference count is not exactly one", failures)
    require(text.count("UNSW Business School (2026)") == 1, "UNSW dataset reference count is not exactly one", failures)
    ref_indices = [refs_text.find(ref) for ref in REFERENCES]
    require(all(i >= 0 for i in ref_indices) and ref_indices == sorted(ref_indices), "references are not alphabetical/in approved order", failures)
    for p in doc.paragraphs:
        if any(p.text.startswith(ref) for ref in REFERENCES):
            require("\n" not in p.text and "\t" not in p.text, f"manual line break or tab in reference: {p.text}", failures)
            require(p.paragraph_format.left_indent and abs(p.paragraph_format.left_indent.cm - 0.75) < 0.05, "reference hanging indent missing", failures)
            require(p.paragraph_format.first_line_indent and abs(p.paragraph_format.first_line_indent.cm + 0.75) < 0.05, "reference first-line indent missing", failures)


def validate_tables_and_artifacts(text: str, doc: Document, failures: list[str]) -> None:
    perf = pd.read_csv(TABLES / "performance_metrics.csv")
    sentiment = pd.read_csv(DATA / "sector_sentiment_index.csv")
    validation = pd.read_csv(TABLES / "pipeline_validation.csv")
    holdings_rows = []
    for table in doc.tables:
        header = [cell.text for cell in table.rows[0].cells]
        if header == ["Fund", "Asset", "Asset class", "Sector", "Target weight", "Final rebalance date"]:
            holdings_rows.extend([[cell.text for cell in row.cells] for row in table.rows[1:]])
        if header[:3] == ["Family", "Calendar", "Annualisation"]:
            require(len(table.rows) == 4, "Table 1 does not contain exactly three data rows", failures)
        if header == ["Fund", "First live date", "End date", "Net annualised return", "Net annualised volatility", "Net Sharpe", "Maximum drawdown"]:
            require(len(table.rows) == 12, "Appendix Table A3 does not contain eleven funds", failures)
    require(perf["fund_id"].nunique() == 11 and len(perf) == 11, "performance metrics do not contain 11 funds", failures)
    require(sentiment["sector"].nunique() == 10, "sentiment artifact does not contain 10 sectors", failures)
    require(validation["status"].eq("PASS").all(), "pipeline validation contains non-PASS rows", failures)
    require((perf.loc[perf["asset_family"].eq("Crypto"), "annualisation_factor"] == 365).all(), "crypto annualisation not 365", failures)
    require((perf.loc[perf["asset_family"].isin(["Equity", "Combined"]), "annualisation_factor"] == 252).all(), "equity/combined annualisation not 252", failures)
    figure1_key_ok(failures)
    for fund in FUNDS:
        require(fund in text, f"fund missing from report: {fund}", failures)
    require(len(holdings_rows) == 38, f"Appendix Table A2 expected 38 positive holdings, observed {len(holdings_rows)}", failures)
    holding_assets = {row[1] for row in holdings_rows}
    require("EOS-USD" not in holding_assets and "ETC-USD" not in holding_assets, "zero-weight crypto asset shown as holding", failures)
    require(all(row[4] != "0.00%" for row in holdings_rows), "zero displayed target weight in A2", failures)


def validate_hashes(failures: list[str]) -> None:
    for rel, expected in EXPECTED_HASHES.items():
        observed = hashlib.sha256((ROOT / rel).read_bytes()).hexdigest()
        require(observed == expected, f"analytical hash changed for {rel}", failures)


def main() -> None:
    failures: list[str] = []
    require(FINAL_MD.exists() and FINAL_MD.stat().st_size > 0, "report_final.md missing or empty", failures)
    require(DOCX.exists() and DOCX.stat().st_size > 0, "report.docx missing or empty", failures)
    require(PDF.exists() and PDF.stat().st_size > 0, "report.pdf missing or empty", failures)
    if failures:
        raise SystemExit("\n".join(failures))

    doc = Document(DOCX)
    text = document_text(doc)
    md = FINAL_MD.read_text(encoding="utf-8")
    combined_text = text + "\n" + md
    stats = word_page_stats()

    forbidden = [
        "DRAFT",
        "STUDENT REVIEW",
        "INSERT TABLE",
        "INSERT FIGURE",
        "REFERENCE DETAILS",
        "TODO",
        "FIXME",
        "placeholder",
        "current holdings",
        "strongest headline",
        "investable in the backtest sense",
        "vader_compound_21d_trailing_lag1",
        "C:\\Users\\",
        "z" + "5641844" + "_projectA",
        "UNSW FINS3645 (2026a)",
        "UNSW FINS3645 (2026b)",
        "(UNSW FINS3645, 2026a)",
        "(UNSW FINS3645, 2026b)",
    ]
    snake_case = [
        "annualised_return_net",
        "annualised_volatility_net",
        "maximum_drawdown_net",
        "total_transaction_cost",
        "average_rebalance_turnover",
        "coverage_quality",
    ]
    require(doc.core_properties.title == TITLE, "DOCX title property is incorrect", failures)
    require(doc.core_properties.author == AUTHOR, "DOCX author property is incorrect", failures)
    require(AUTHOR in text, "author not present in report text", failures)
    require(ZID in text, "zID not present in report text", failures)
    require(stats["refs"] - stats["exec"] <= 10, f"narrative exceeds 10 pages: {stats['refs'] - stats['exec']}", failures)
    require(stats["exec"] == 2, "Executive Summary does not begin on physical page 2", failures)
    require(stats["refs"] > stats["conclusion"], "References do not begin after Conclusion", failures)
    require(stats["last_narr_chars"] >= 900, f"last narrative page appears almost empty: {stats['last_narr_chars']} characters", failures)
    require(count_page_fields(doc) <= 2, "duplicate PAGE fields detected in footers", failures)
    section_page_numbering_ok(doc, failures)
    for term in forbidden:
        require(term not in combined_text, f"forbidden term present: {term}", failures)
    for term in snake_case:
        require(term not in combined_text, f"snake_case analytical field present: {term}", failures)
    require("*" not in text, "literal Markdown asterisk appears in DOCX text", failures)
    require(len(re.findall(r"Recommendation [123] is to", text)) == 3, "report does not contain exactly three recommendations", failures)
    require(
        "The highest measured net annualised return and net Sharpe in this sample is" not in combined_text,
        "incorrect Executive Summary grammar phrase remains",
        failures,
    )
    require(
        "PASS - deterministic CSV hashes matched across two complete pipeline runs" in text,
        "Appendix A6 idempotence PASS wording is missing",
        failures,
    )
    for i in range(1, 7):
        require(f"Figure {i}." in text, f"Figure {i} caption missing", failures)
    for i in range(1, 4):
        require(f"Table {i}." in text, f"Table {i} caption missing", failures)
    require(media_count(DOCX) >= 7, "expected at least seven embedded images", failures)
    for value in ["49.25%", "0.91", "76.12%", "-73.55%", "49.57%", "48.85%", "146,830", "146,836", "10,060", "-0.0705", "93.61%", "81 PASS"]:
        require(value in text, f"missing factual value {value}", failures)
    validate_references(text, doc, failures)
    validate_tables_and_artifacts(text, doc, failures)
    validate_hashes(failures)

    pdf_pages = len(re.findall(rb"/Type\s*/Page\b", PDF.read_bytes()))
    require(pdf_pages == stats["pages"], f"PDF page count {pdf_pages} differs from Word page count {stats['pages']}", failures)

    if failures:
        raise SystemExit("report validation failed:\n" + "\n".join(failures))
    print("report validation passed")
    print(f"docx_size={DOCX.stat().st_size}")
    print(f"pdf_size={PDF.stat().st_size}")
    print(f"word_pages={stats['pages']}")
    print(f"narrative_pages={stats['refs'] - stats['exec']}")
    print(f"pdf_pages={pdf_pages}")
    print(f"embedded_images={media_count(DOCX)}")
    print(f"tables={len(doc.tables)}")


if __name__ == "__main__":
    main()
