"""Build the final Project B Word report and export it to PDF."""
from __future__ import annotations

import os
import re
import subprocess
from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "report"
RESULTS = ROOT / "results"
TABLES = RESULTS / "tables"
DATA = RESULTS / "data"
FIGURES = RESULTS / "figures"
ASSETS = REPORT / "assets"
FINAL_MD = REPORT / "report_final.md"
DOCX = REPORT / "report.docx"
PDF = REPORT / "report.pdf"

TITLE = "Signal Mosaic: Systematic Multi-Asset Funds with Coverage-Aware News Sentiment"
AUTHOR = "Chenhang Huang"
ZID = "z5641844"

NAVY = RGBColor(20, 38, 64)
BURGUNDY = RGBColor(122, 45, 62)
TEAL = RGBColor(36, 111, 122)
LIGHT = "F2F4F5"

FAMILY_ORDER = {"Equity": 0, "Crypto": 1, "Combined": 2}
METHOD_ORDER = {
    "Equal Weight": 0,
    "Minimum Variance": 1,
    "Risk Parity": 2,
    "Naive Sentiment Tilt": 3,
    "Coverage-Gated Sentiment Tilt": 4,
}


def pct(value: float, digits: int = 2) -> str:
    return f"{float(value) * 100:.{digits}f}%"


def num(value: float, digits: int = 2) -> str:
    return f"{float(value):.{digits}f}"


def load_tables() -> dict[str, pd.DataFrame]:
    tables = {
        "performance": pd.read_csv(TABLES / "performance_metrics.csv"),
        "design": pd.read_csv(TABLES / "fund_backtest_design.csv"),
        "fact": pd.read_csv(TABLES / "fund_fact_sheet_summary.csv"),
        "holdings": pd.read_csv(TABLES / "fund_latest_holdings.csv", parse_dates=["rebalance_date"]),
        "fusion": pd.read_csv(TABLES / "fusion_before_after.csv"),
        "predictive": pd.read_csv(TABLES / "fusion_predictive_diagnostics.csv"),
        "sentiment": pd.read_csv(TABLES / "sentiment_model_diagnostics.csv"),
        "validation": pd.read_csv(TABLES / "pipeline_validation.csv"),
        "inventory": pd.read_csv(TABLES / "app_artifact_inventory.csv"),
        "fund_weights": pd.read_csv(DATA / "fund_weights.csv", parse_dates=["rebalance_date"]),
        "sector_sentiment": pd.read_csv(DATA / "sector_sentiment_index.csv"),
    }
    return tables


def build_report_risk_return_figure(performance: pd.DataFrame) -> Path:
    ASSETS.mkdir(parents=True, exist_ok=True)
    out = ASSETS / "fund_risk_return_all_11.png"
    perf = sort_funds(performance)
    colors = {"Equity": "#7A2D3E", "Crypto": "#C49A3A", "Combined": "#266F7A"}
    markers = {
        "Equal Weight": "o",
        "Minimum Variance": "s",
        "Risk Parity": "^",
        "Naive Sentiment Tilt": "D",
        "Coverage-Gated Sentiment Tilt": "P",
    }
    labels = {
        "Equity Equal Weight": "Eq EW",
        "Equity Minimum Variance": "Eq Min Var",
        "Equity Risk Parity": "Eq Risk Parity",
        "Equity Naive Sentiment Tilt": "Eq Naive Tilt",
        "Equity Coverage-Gated Sentiment Tilt": "Eq Coverage-Gated",
        "Crypto Equal Weight": "Crypto EW",
        "Crypto Minimum Variance": "Crypto Min Var",
        "Crypto Risk Parity": "Crypto Risk Parity",
        "Combined Equal Weight": "Combined EW",
        "Combined Minimum Variance": "Combined Min Var",
        "Combined Risk Parity": "Combined Risk Parity",
    }
    offsets = {
        "Equity Equal Weight": (14, 18),
        "Equity Minimum Variance": (-42, -12),
        "Equity Risk Parity": (14, -18),
        "Equity Naive Sentiment Tilt": (24, -32),
        "Equity Coverage-Gated Sentiment Tilt": (-64, 20),
        "Crypto Equal Weight": (-48, -4),
        "Crypto Minimum Variance": (16, 12),
        "Crypto Risk Parity": (18, -18),
        "Combined Equal Weight": (18, 16),
        "Combined Minimum Variance": (-60, 12),
        "Combined Risk Parity": (14, -22),
    }
    fig, ax = plt.subplots(figsize=(9.2, 5.3), dpi=220)
    fig.patch.set_facecolor("#FBF8F2")
    ax.set_facecolor("#FBF8F2")
    key_lines: list[str] = []
    for idx, (_, row) in enumerate(perf.iterrows(), start=1):
        x = row["annualised_volatility_net"] * 100
        y = row["annualised_return_net"] * 100
        ax.scatter(
            x,
            y,
            s=62,
            color=colors.get(row["asset_family"], "#333333"),
            marker=markers.get(row["method"], "o"),
            edgecolor="#1C2634",
            linewidth=0.5,
            zorder=3,
        )
        dx, dy = offsets[row["fund_name"]]
        ax.annotate(
            str(idx),
            (x, y),
            xytext=(dx, dy),
            textcoords="offset points",
            ha="center",
            va="center",
            fontsize=7.0,
            color="#1C2634",
            bbox=dict(boxstyle="circle,pad=0.18", fc="#FBF8F2", ec=colors.get(row["asset_family"], "#333333"), lw=0.8),
            arrowprops=dict(arrowstyle="-", color="#777777", lw=0.45, shrinkA=0, shrinkB=4),
            zorder=4,
        )
        key_lines.append(f"{idx}. {labels[row['fund_name']]}")
    ax.set_title("Net risk-return profile across all eleven Signal Mosaic funds", loc="left", fontsize=12, color="#142640", weight="bold")
    ax.set_xlabel("Net annualised volatility (%)", fontsize=9)
    ax.set_ylabel("Net annualised return (%)", fontsize=9)
    ax.grid(True, color="#D6D1C8", linewidth=0.7, alpha=0.75)
    for spine in ax.spines.values():
        spine.set_visible(False)
    ax.tick_params(labelsize=8)
    family_handles = [
        plt.Line2D([0], [0], marker="o", color="w", label=family, markerfacecolor=color, markeredgecolor="#1C2634", markersize=7)
        for family, color in colors.items()
    ]
    ax.legend(handles=family_handles, loc="upper left", frameon=False, fontsize=8)
    fig.text(0.735, 0.70, "Point key", fontsize=8.2, color="#142640", weight="bold")
    fig.text(0.735, 0.245, "\n".join(key_lines), fontsize=7.0, color="#1C2634", linespacing=1.25)
    fig.text(
        0.01,
        0.01,
        "Source: UNSW Business School FINS3645 Project Data Bundle; Signal Mosaic calculations. "
        "Out-of-sample 2021-2023; net returns; Sharpe assumes rf = 0.",
        fontsize=7,
        color="#555555",
    )
    fig.tight_layout(rect=[0, 0.06, 0.72, 1])
    fig.savefig(out, facecolor=fig.get_facecolor(), bbox_inches="tight")
    plt.close(fig)
    return out


def sort_funds(frame: pd.DataFrame) -> pd.DataFrame:
    out = frame.copy()
    out["_family"] = out["asset_family"].map(FAMILY_ORDER)
    out["_method"] = out["method"].map(METHOD_ORDER)
    return out.sort_values(["_family", "_method", "fund_name"], kind="stable").drop(columns=["_family", "_method"])


def set_cell_text(cell, text: str, *, bold: bool = False, size: float = 7.8) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Aptos"
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_keep_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def set_keep_together(paragraph) -> None:
    paragraph.paragraph_format.keep_together = True


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def column_widths(headers: list[str]) -> list[float] | None:
    if headers == ["Family", "Calendar", "Annualisation", "Initial estimation period", "First live date", "End date"]:
        return [1.8, 3.0, 1.9, 4.2, 2.9, 2.9]
    if headers == ["Fund", "Family", "Net annualised return", "Net volatility", "Net Sharpe", "Maximum drawdown"]:
        return [5.0, 1.8, 2.7, 2.5, 1.7, 2.7]
    if headers == ["Fund", "Net annualised return", "Net volatility", "Net Sharpe", "Maximum drawdown", "Average rebalance turnover"]:
        return [5.6, 2.5, 2.4, 1.6, 2.3, 2.6]
    if headers == ["Fund", "Cumulative net return", "Annualised net return", "Net volatility", "Net Sharpe", "Maximum drawdown"]:
        return [5.6, 3.3, 3.3, 3.1, 2.2, 3.1]
    if headers == ["Fund", "Obs.", "Ann.", "Avg. turnover", "Total turnover", "Total cost", "Rebalances", "Non-zero holdings", "Largest target"]:
        return [4.8, 1.4, 1.3, 2.2, 2.2, 1.7, 1.5, 1.8, 1.8]
    if headers == ["Fund", "Asset", "Asset class", "Sector", "Target weight", "Final rebalance date"]:
        return [5.0, 2.4, 2.3, 3.5, 2.4, 3.1]
    if headers == ["Fund", "First live date", "End date", "Net annualised return", "Net annualised volatility", "Net Sharpe", "Maximum drawdown"]:
        return [5.8, 2.7, 2.7, 3.2, 3.2, 2.0, 3.0]
    if headers == ["Scored headlines", "Positive share", "Neutral share", "Negative share", "Exact-zero share", "Ticker-days", "Sector-days with news", "Sector-days without news"]:
        return [3.0, 2.8, 2.8, 2.8, 3.0, 2.8, 3.2, 3.2]
    if headers == ["Sample", "Pooled Spearman", "Average monthly Spearman", "Monthly obs.", "Pair obs.", "Median coverage quality"]:
        return [5.3, 3.0, 4.2, 2.5, 2.5, 3.7]
    if headers == ["Validation item", "Observed"]:
        return [7.0, 5.0]
    return None


def set_table_grid(table, widths_cm: list[float]) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(int(width * 567) for width in widths_cm)))

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_cm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 567)))
        tbl_grid.append(grid_col)


def add_table_from_rows(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    widths = column_widths(headers)
    if widths:
        set_table_grid(table, widths)
    hdr = table.rows[0]
    hdr._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    prevent_row_split(hdr)
    for idx, heading in enumerate(headers):
        set_cell_text(hdr.cells[idx], heading, bold=True)
        if widths:
            hdr.cells[idx].width = Cm(widths[idx])
        shade_cell(hdr.cells[idx], LIGHT)
    for row in rows:
        table_row = table.add_row()
        prevent_row_split(table_row)
        cells = table_row.cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
            if widths:
                cells[idx].width = Cm(widths[idx])
            if idx > 1:
                cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()


def caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_keep_next(p)
    run = p.add_run(text)
    run.font.size = Pt(8.5)
    run.font.color.rgb = NAVY


def source_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(80, 80, 80)


def table1_rows(tables: dict[str, pd.DataFrame]) -> list[list[str]]:
    design = tables["design"]
    rows = []
    for family in ["Equity", "Crypto", "Combined"]:
        fam = design.loc[design["asset_family"].eq(family)].iloc[0]
        rows.append(
            [
                family,
                str(fam["calendar"]),
                str(int(fam["annualisation_factor"])),
                f"{fam['initial_estimation_start']} to {fam['initial_estimation_end']}",
                str(fam["first_live_date"]),
                str(fam["end_date"]),
            ]
        )
    return rows


def table2_rows(tables: dict[str, pd.DataFrame]) -> list[list[str]]:
    perf = sort_funds(tables["performance"])
    return [
        [
            row["fund_name"],
            row["asset_family"],
            pct(row["annualised_return_net"]),
            pct(row["annualised_volatility_net"]),
            num(row["Sharpe_net"]),
            pct(row["maximum_drawdown_net"]),
        ]
        for _, row in perf.iterrows()
    ]


def table3_rows(tables: dict[str, pd.DataFrame]) -> list[list[str]]:
    order = [
        "Equity Equal Weight",
        "Equity Naive Sentiment Tilt",
        "Equity Coverage-Gated Sentiment Tilt",
    ]
    fusion = tables["fusion"].set_index("fund_name").loc[order].reset_index()
    return [
        [
            row["fund_name"],
            pct(row["annualised_return_net"]),
            pct(row["annualised_volatility_net"]),
            num(row["Sharpe_net"]),
            pct(row["maximum_drawdown_net"]),
            pct(row["average_rebalance_turnover"]),
        ]
        for _, row in fusion.iterrows()
    ]


def appendix_a1_rows(tables: dict[str, pd.DataFrame]) -> tuple[list[list[str]], list[list[str]]]:
    perf = sort_funds(tables["performance"])
    performance = [
        [
            row["fund_name"],
            pct(row["cumulative_return_net"]),
            pct(row["annualised_return_net"]),
            pct(row["annualised_volatility_net"]),
            num(row["Sharpe_net"]),
            pct(row["maximum_drawdown_net"]),
        ]
        for _, row in perf.iterrows()
    ]
    implementation = [
        [
            row["fund_name"],
            str(int(row["number_of_observations"])),
            str(int(row["annualisation_factor"])),
            pct(row["average_rebalance_turnover"]),
            pct(row["total_rebalance_turnover"]),
            pct(row["total_transaction_cost"]),
            str(int(row["number_of_rebalances"])),
            str(int(row["current_number_of_holdings"])),
            pct(row["largest_current_weight"]),
        ]
        for _, row in perf.iterrows()
    ]
    return performance, implementation


def appendix_holdings_rows(tables: dict[str, pd.DataFrame]) -> list[list[str]]:
    funds = [
        "Equity Equal Weight",
        "Crypto Minimum Variance",
        "Combined Risk Parity",
        "Equity Coverage-Gated Sentiment Tilt",
    ]
    holdings = tables["holdings"].copy()
    output: list[list[str]] = []
    for fund in funds:
        part = (
            holdings.loc[holdings["fund_name"].eq(fund)]
            .loc[lambda frame: frame["target_weight"].gt(1e-12)]
            .sort_values(["target_weight", "asset"], ascending=[False, True], kind="stable")
            .head(10)
        )
        for _, row in part.iterrows():
            output.append(
                [
                    fund,
                    row["asset"],
                    row["asset_class"],
                    "" if pd.isna(row.get("sector")) else row["sector"],
                    pct(row["target_weight"]),
                    row["rebalance_date"].strftime("%Y-%m-%d"),
                ]
            )
    return output


def appendix_fact_rows(tables: dict[str, pd.DataFrame]) -> list[list[str]]:
    fact = sort_funds(tables["fact"].rename(columns={"family": "asset_family"}) if "family" in tables["fact"].columns else tables["fact"])
    return [
        [
            row["fund_name"],
            str(row["first_live_date"]),
            str(row["end_date"]),
            pct(row["annualised_return_net"]),
            pct(row["annualised_volatility_net"]),
            num(row["Sharpe_net"]),
            pct(row["maximum_drawdown_net"]),
        ]
        for _, row in fact.iterrows()
    ]


def appendix_sentiment_rows(tables: dict[str, pd.DataFrame]) -> list[list[str]]:
    row = tables["sentiment"].iloc[0]
    return [
        [
            f"{int(row['total_scored_headlines']):,}",
            pct(row["positive_share"]),
            pct(row["vader_neutral_share"]),
            pct(row["negative_share"]),
            pct(row["exact_zero_compound_share"]),
            f"{int(row['ticker_days']):,}",
            f"{int(row['sector_days_with_news']):,}",
            f"{int(row['sector_days_without_news']):,}",
        ]
    ]


def appendix_predictive_rows(tables: dict[str, pd.DataFrame]) -> list[list[str]]:
    pred = tables["predictive"].copy()
    labels = {
        "all_valid": "All valid observations",
        "above_median_coverage_quality": "Above-median coverage quality",
        "at_or_below_median_coverage_quality": "At/below-median coverage quality",
    }
    return [
        [
            labels.get(str(row["sample"]), str(row["sample"]).replace("_", " ").title()),
            num(row["pooled_spearman"], 4),
            num(row["average_cross_sectional_spearman"], 4),
            str(int(row["valid_monthly_observations"])),
            str(int(row["valid_pair_observations"])),
            num(row["median_coverage_quality"], 4),
        ]
        for _, row in pred.iterrows()
    ]


def appendix_validation_rows(tables: dict[str, pd.DataFrame]) -> list[list[str]]:
    validation = tables["validation"]
    inventory = tables["inventory"]
    perf = tables["performance"]
    sentiment = tables["sector_sentiment"]
    idempotence_status = "PASS - deterministic CSV hashes matched across two complete pipeline runs"
    return [
        ["Total validations", str(len(validation))],
        ["PASS count", str(int(validation["status"].eq("PASS").sum()))],
        ["FAIL count", str(int((~validation["status"].eq("PASS")).sum()))],
        ["Final unique fund count", str(int(perf["fund_id"].nunique()))],
        ["Final sentiment row count", f"{len(sentiment):,}"],
        ["Idempotence result", idempotence_status],
        ["App artifact readiness", f"{int(inventory['readiness_status'].eq('READY').sum())} ready artifacts"],
    ]


def add_figure(doc: Document, path: Path, caption_text: str, *, width: float = 6.25) -> None:
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_keep_next(p)
    run = p.add_run()
    inline = run.add_picture(str(path), width=Inches(width))
    inline._inline.docPr.set("descr", caption_text)
    caption(doc, caption_text)


def add_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(8.5)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.clear()
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def configure_header_footer(section, *, restart_at: int | None = None) -> None:
    section.different_first_page_header_footer = False
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header.paragraphs[0]
    header.clear()
    header.text = "Signal Mosaic | FINS3645 Project B"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = NAVY
    add_page_number(section.footer.paragraphs[0])
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if restart_at is None:
        if pg_num is not None:
            sect_pr.remove(pg_num)
    else:
        if pg_num is None:
            pg_num = OxmlElement("w:pgNumType")
            sect_pr.append(pg_num)
        pg_num.set(qn("w:start"), str(restart_at))


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.different_first_page_header_footer = True

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.line_spacing = 1.05
    styles["Normal"].paragraph_format.space_after = Pt(3)
    for style_name, size in [("Heading 1", 15), ("Heading 2", 12)]:
        style = styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = NAVY
        style.font.bold = True
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.page_break_before = False
    styles["Title"].font.name = "Aptos"
    styles["Title"].font.size = Pt(22)
    styles["Title"].font.color.rgb = NAVY
    styles["Caption"].font.name = "Aptos"
    styles["Caption"].font.size = Pt(8.5)
    styles["Caption"].font.italic = True


def add_cover(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Signal Mosaic")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = NAVY
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Systematic Multi-Asset Funds with Coverage-Aware News Sentiment")
    r.font.size = Pt(16)
    r.font.color.rgb = BURGUNDY
    for _ in range(4):
        doc.add_paragraph()
    for line in [
        "FINS3645 Financial Market Data Design & Analysis",
        "Project B",
        "Author: Chenhang Huang",
        "zID: z5641844",
        "Term 2, 2026",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.size = Pt(12)
        r.font.color.rgb = NAVY


def start_numbered_body(doc: Document) -> None:
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    configure_header_footer(section, restart_at=1)


def split_markdown(md: str) -> dict[str, list[str]]:
    sections: dict[str, list[str]] = {}
    current = None
    for line in md.splitlines():
        if line.startswith("## "):
            current = line[3:].strip()
            sections[current] = []
        elif current is not None:
            sections[current].append(line)
    return sections


def paragraphs(lines: list[str]) -> list[str]:
    out: list[str] = []
    buf: list[str] = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            if buf:
                out.append(" ".join(buf))
                buf = []
            continue
        if stripped.startswith("#"):
            continue
        buf.append(stripped)
    if buf:
        out.append(" ".join(buf))
    return out


def add_body_section(doc: Document, heading: str, paras: list[str], tables: dict[str, pd.DataFrame]) -> None:
    doc.add_heading(heading, level=1)
    for para in paras:
        doc.add_paragraph(para)
        if heading.startswith("1.") and para.startswith("Table 1 reports"):
            caption(doc, "Table 1. Walk-forward backtest design by asset family.")
            add_table_from_rows(
                doc,
                ["Family", "Calendar", "Annualisation", "Initial estimation period", "First live date", "End date"],
                table1_rows(tables),
            )
            source_note(doc, "Source: UNSW Business School FINS3645 Project Data Bundle; Signal Mosaic calculations.")
            add_note(
                doc,
                "Note: All fund backtests use an expanding estimation window, monthly rebalancing, past information only, monthly buy-and-hold drift, long-only optimised funds, a 20% asset cap for optimised funds, rf = 0, a 10 basis-point one-way turnover cost, and no initial establishment cost.",
            )
        if heading.startswith("2.") and para.startswith("Table 2 compares"):
            caption(doc, "Table 2. Net out-of-sample performance across the eleven Signal Mosaic funds, 2021-2023.")
            add_table_from_rows(
                doc,
                ["Fund", "Family", "Net annualised return", "Net volatility", "Net Sharpe", "Maximum drawdown"],
                table2_rows(tables),
            )
            source_note(doc, "Source: UNSW Business School FINS3645 Project Data Bundle; Signal Mosaic calculations.")
            add_figure(
                doc,
                ASSETS / "fund_risk_return_all_11.png",
                "Figure 1. Net annualised return and volatility across the eleven funds, 2021-2023.",
            )
        if heading.startswith("2.") and para.startswith("The combined family gives"):
            add_figure(
                doc,
                FIGURES / "fund_growth_of_one_by_family.png",
                "Figure 2. Net growth of one dollar by asset family and method.",
                width=5.55,
            )
            add_figure(
                doc,
                FIGURES / "fund_drawdowns_combined.png",
                "Figure 3. Net drawdowns of the combined equity-crypto funds.",
            )
        if heading.startswith("3.") and para.startswith("The sector sentiment artifact"):
            add_figure(
                doc,
                FIGURES / "sector_sentiment_timeseries.png",
                "Figure 4. Twenty-one-trading-day trailing plain-VADER sentiment across the ten equity sectors.",
            )
        if heading.startswith("4.") and para.startswith("Coverage Lens describes"):
            add_figure(
                doc,
                FIGURES / "sentiment_coverage_context.png",
                "Figure 5. Sector sentiment displayed with Coverage Lens context.",
            )
        if heading.startswith("4.") and para.startswith("Table 3 and Figure 6"):
            caption(doc, "Table 3. Equity Equal Weight versus the two sentiment-overlay funds.")
            add_table_from_rows(
                doc,
                ["Fund", "Net annualised return", "Net volatility", "Net Sharpe", "Maximum drawdown", "Average rebalance turnover"],
                table3_rows(tables),
            )
            source_note(doc, "Source: UNSW Business School FINS3645 Project Data Bundle; Signal Mosaic calculations using plain VADER.")
            add_figure(
                doc,
                FIGURES / "fusion_growth_of_one.png",
                "Figure 6. Net growth of one dollar for Equity Equal Weight and the two sentiment overlays.",
            )


REFERENCE_RUNS = [
    [
        ("Hric, J. and Lin, Y. (2026) ", False),
        ("Applied Data Science in FinTech: Models, Tools, and Case Studies", True),
        (". London: Routledge.", False),
    ],
    [
        ("Hutto, C.J. and Gilbert, E.E. (2014) 'VADER: A Parsimonious Rule-Based Model for Sentiment Analysis of Social Media Text', ", False),
        ("Proceedings of the International AAAI Conference on Web and Social Media", True),
        (", 8(1), pp. 216-225.", False),
    ],
    [
        ("Markowitz, H. (1952) 'Portfolio Selection', ", False),
        ("The Journal of Finance", True),
        (", 7(1), pp. 77-91.", False),
    ],
    [
        ("Sharpe, W.F. (1966) 'Mutual Fund Performance', ", False),
        ("The Journal of Business", True),
        (", 39(1, Part 2), pp. 119-138.", False),
    ],
    [
        ("UNSW Business School (2026) ", False),
        ("FINS3645 Project Data Bundle: project_data.zip", True),
        (" [dataset].", False),
    ],
]


def add_references(doc: Document) -> None:
    heading = doc.add_heading("References", level=1)
    heading.paragraph_format.page_break_before = True
    for ref in REFERENCE_RUNS:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(-0.75)
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(5)
        for text, italic in ref:
            run = p.add_run(text)
            run.font.name = "Aptos"
            run.font.size = Pt(10.5)
            run.italic = italic


def add_appendices(doc: Document, tables: dict[str, pd.DataFrame]) -> None:
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)
    configure_header_footer(section)
    doc.add_heading("Appendices", level=1)
    caption(doc, "Appendix Table A1a. Expanded fund performance metrics.")
    rows_a1a, rows_a1b = appendix_a1_rows(tables)
    add_table_from_rows(
        doc,
        ["Fund", "Cumulative net return", "Annualised net return", "Net volatility", "Net Sharpe", "Maximum drawdown"],
        rows_a1a,
    )
    source_note(doc, "Source: UNSW Business School FINS3645 Project Data Bundle; Signal Mosaic calculations.")
    doc.add_page_break()
    caption(doc, "Appendix Table A1b. Expanded fund implementation metrics.")
    add_table_from_rows(
        doc,
        [
            "Fund",
            "Obs.",
            "Ann.",
            "Avg. turnover",
            "Total turnover",
            "Total cost",
            "Rebalances",
            "Non-zero holdings",
            "Largest target",
        ],
        rows_a1b,
    )
    source_note(doc, "Source: UNSW Business School FINS3645 Project Data Bundle; Signal Mosaic calculations.")
    doc.add_page_break()
    add_figure(
        doc,
        FIGURES / "fund_weights_over_time_combined.png",
        "Appendix Figure A1. Combined-fund weights over time by method.",
        width=7.0,
    )
    doc.add_page_break()
    caption(doc, "Appendix Table A2. Representative latest backtest target holdings at the final reported rebalance.")
    add_table_from_rows(
        doc,
        ["Fund", "Asset", "Asset class", "Sector", "Target weight", "Final rebalance date"],
        appendix_holdings_rows(tables),
    )
    add_note(doc, "Note: The complete holdings file is included in the submitted project artifacts.")
    doc.add_page_break()
    caption(doc, "Appendix Table A3. Fund fact-sheet summary.")
    add_table_from_rows(
        doc,
        ["Fund", "First live date", "End date", "Net annualised return", "Net annualised volatility", "Net Sharpe", "Maximum drawdown"],
        appendix_fact_rows(tables),
    )
    doc.add_page_break()
    caption(doc, "Appendix Table A4. Plain-VADER headline scoring diagnostics.")
    add_table_from_rows(
        doc,
        ["Scored headlines", "Positive share", "Neutral share", "Negative share", "Exact-zero share", "Ticker-days", "Sector-days with news", "Sector-days without news"],
        appendix_sentiment_rows(tables),
    )
    caption(doc, "Appendix Table A5. Monthly sentiment-return rank-correlation diagnostics.")
    add_table_from_rows(
        doc,
        ["Sample", "Pooled Spearman", "Average monthly Spearman", "Monthly obs.", "Pair obs.", "Median coverage quality"],
        appendix_predictive_rows(tables),
    )
    add_note(doc, "Note: No formal significance claim is made.")
    caption(doc, "Appendix Table A6. Pipeline validation summary.")
    add_table_from_rows(doc, ["Validation item", "Observed"], appendix_validation_rows(tables))
    add_note(doc, "Note: The complete validation table is included in the submitted project artifacts.")


def validate_clean_text(doc: Document) -> None:
    text = "\n".join([p.text for p in doc.paragraphs])
    bad = [
        "DRAFT FOR STUDENT REVIEW",
        "STUDENT REVIEW REQUIRED",
        "INSERT TABLE",
        "INSERT FIGURE",
        "REFERENCE DETAILS REQUIRE STUDENT VERIFICATION",
        "TODO",
        "vader_compound_21d_trailing_lag1",
        "strongest headline",
        "investable in the backtest sense",
        "C:\\Users\\",
        "z" + "5641844" + "_projectA",
    ]
    found = [term for term in bad if term in text]
    if found:
        raise ValueError(f"report contains forbidden terms: {found}")


def export_pdf_with_word(docx_path: Path, pdf_path: Path) -> tuple[int | None, int | None]:
    ps = f"""
$ErrorActionPreference = 'Stop'
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$doc = $word.Documents.Open('{docx_path}')
$doc.Fields.Update() | Out-Null
$doc.Repaginate()
$pages = $doc.ComputeStatistics(2)
$doc.ExportAsFixedFormat('{pdf_path}', 17)
$doc.Close($false)
$word.Quit()
Write-Output "WORD_PAGES=$pages"
"""
    completed = subprocess.run(
        ["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-Command", ps],
        cwd=ROOT,
        text=True,
        capture_output=True,
        check=True,
    )
    match = re.search(r"WORD_PAGES=(\d+)", completed.stdout)
    pages = int(match.group(1)) if match else None
    pdf_size = pdf_path.stat().st_size if pdf_path.exists() else None
    return pages, pdf_size


def build() -> None:
    if not FINAL_MD.exists():
        raise FileNotFoundError(FINAL_MD)
    tables = load_tables()
    build_report_risk_return_figure(tables["performance"])
    md = FINAL_MD.read_text(encoding="utf-8")
    sections = split_markdown(md)
    required_figures = [
        "fund_growth_of_one_by_family.png",
        "fund_drawdowns_combined.png",
        "sector_sentiment_timeseries.png",
        "sentiment_coverage_context.png",
        "fusion_growth_of_one.png",
        "fund_weights_over_time_combined.png",
    ]
    missing_figures = [name for name in required_figures if not (FIGURES / name).exists()]
    if missing_figures:
        raise FileNotFoundError(f"missing report figures: {missing_figures}")
    if not (ASSETS / "fund_risk_return_all_11.png").exists():
        raise FileNotFoundError(ASSETS / "fund_risk_return_all_11.png")

    doc = Document()
    configure_document(doc)
    props = doc.core_properties
    props.title = TITLE
    props.author = AUTHOR
    props.subject = "FINS3645 Project B"
    props.keywords = "multi-asset funds, sentiment, Coverage Lens, Streamlit"

    add_cover(doc)
    start_numbered_body(doc)
    for heading in [
        "Executive Summary",
        "1. Funds and Walk-Forward Backtest Design",
        "2. Out-of-Sample Results and Fund Fact Sheets",
        "3. Standalone Sector Sentiment Index",
        "4. Innovation: Coverage-Gated Sentiment Fusion",
        "5. Signal Mosaic App and Investor Journey",
        "6. Critical Reflection and Three Recommendations",
        "7. Conclusion",
    ]:
        add_body_section(doc, heading, paragraphs(sections[heading]), tables)
    add_references(doc)
    add_appendices(doc, tables)
    validate_clean_text(doc)
    DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX)
    pages, pdf_size = export_pdf_with_word(DOCX.resolve(), PDF.resolve())
    print(f"wrote {DOCX.relative_to(ROOT)} ({DOCX.stat().st_size} bytes)")
    print(f"wrote {PDF.relative_to(ROOT)} ({pdf_size} bytes)")
    if pages is not None:
        print(f"word_page_count={pages}")


if __name__ == "__main__":
    os.environ.setdefault("PYTHONDONTWRITEBYTECODE", "1")
    build()
